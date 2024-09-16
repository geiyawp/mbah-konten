import streamlit as st
from openai import OpenAI
import requests
from bs4 import BeautifulSoup
from io import BytesIO
from docx import Document
from docx.shared import Inches  # Used for setting image size

def generate_titles(client, topic, language):
    if language == "Bahasa Indonesia":
        prompt = f"Generate 10 compelling and unique titles for content about {topic}, in Bahasa Indonesia."
    else:
        prompt = f"Generate 10 compelling and unique titles for content about {topic}, in English"
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": f"You are a creative title generator. Your task is to generate a list of 10 potential titles ONLY SEPARATED BY COMMAS WITHOUT OTHER FORMATTING. DO NOT OUTPUT COMMENT OR ANYTHING ELSE."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=150
    )
    titles = response.choices[0].message.content.strip().split(',')
    return titles

def suggest_keywords(client, title, language):
    if language == "Bahasa Indonesia":
        prompt = f"Generate 1-5 important SEO keywords for an article titled '{title}', in Bahasa Indonesia."
    else:
        prompt = f"Generate 1-5 important SEO keywords for an article titled '{title}', in English."
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": f"You are an expert in identifying SEO keywords. Your task is to generate a list of potential keywords ONLY SEPARATED BY COMMAS WITHOUT OTHER FORMATTING. DO NOT OUTPUT COMMENT OR ANYTHING ELSE."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=150
    )
    keywords = response.choices[0].message.content.strip().split(',')
    return [kw.strip() for kw in keywords]

def generate_image(client, title, size):
    prompt = f"Create a detailed artistic image representation for a post titled '{title}'."
    response = client.images.generate(
        prompt=prompt,
        model="dall-e-3",
        size=size
    )
    return response.data[0].url

def generate_article(client, content_type, title, keywords, tone, word_count, reference_texts, language):
    if language == "Bahasa Indonesia":
        prompt = (f"Menggunakan referensi: {reference_texts}. Buat {content_type} yang ramah SEO dan menarik menggunakan Bahasa Indonesia dengan judul '{title}' "
                  f"dengan kata kunci berikut: {keywords}. Gunakan nada {tone}. Sekitar {word_count} kata.")
    else:
        prompt = (f"Using the references: {reference_texts}. Create a SEO-friendly and engaging {content_type} in English, titled '{title}' "
                  f"with the following keywords: {keywords}. Use a {tone} tone. In approximately {word_count} words.")
        
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": f"You are an expert writer producing engaging content. Your task is to write great content with PROPER FORMATTING ACCORDING TO THE CONTENT TYPE. DO NOT OUTPUT COMMENT OR ANYTHING ELSE"},
            {"role": "user", "content": prompt}
        ],
        max_tokens=4096  # Convert words to approximate token count
    )
    article = response.choices[0].message.content.strip()
    return article

def retrieve_references(urls):
    texts = []
    for url in urls.strip().split('\n'):
        url = url.strip()
        if url:
            try:
                response = requests.get(url)
                soup = BeautifulSoup(response.content, 'html.parser')
                texts.append(soup.get_text())
            except Exception as e:
                st.warning(f"Could not retrieve {url}: {e}")
    return ' '.join(texts[:3])  # Just use a few references for demo

def save_as_docx(title, article, image_url=None):
    doc = Document()
    doc.add_heading(title, 0)
    if image_url:
        image_response = requests.get(image_url)
        if image_response.status_code == 200:
            image_data = BytesIO(image_response.content)
            doc.add_picture(image_data, width=Inches(6))
    doc.add_paragraph(article)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    # Initialize session state variables
    if 'client' not in st.session_state:
        st.session_state.client = None
    if 'titles' not in st.session_state:
        st.session_state.titles = []
    if 'keywords' not in st.session_state:
        st.session_state.keywords = []
    if 'image_url' not in st.session_state:
        st.session_state.image_url = None
    if 'article' not in st.session_state:
        st.session_state.article = None
    if 'image_size' not in st.session_state:
        st.session_state.image_size = None

    if st.session_state.client is None:
        api_key = st.sidebar.text_input("OpenAI API Key", type="password")
        if api_key:
            st.session_state.client = OpenAI(api_key=api_key)
        else:
            st.error("Please provide an OpenAI API key.")
            return

    client = st.session_state.client

    st.title("Mbah Konten v1.0")

    content_type = st.selectbox("Select Content Type", ["Blog Post", "LinkedIn Post"])
    language = st.selectbox("Select Language", ["English", "Bahasa Indonesia"])
    topic = st.text_input("What do you want to write about?")

    if topic and st.button("Generate Titles"):
        st.session_state.titles = generate_titles(client, topic, language)
        st.session_state.keywords = []

    if st.session_state.titles:
        title = st.selectbox("Choose a Title", st.session_state.titles)

        if title:
            if not st.session_state.keywords:
                st.session_state.keywords = suggest_keywords(client, title, language)

            suggested_keywords = st.session_state.keywords
            selected_keywords = st.multiselect("Select Keywords", suggested_keywords, default=suggested_keywords)
            added_keywords = st.text_input("Add Additional Keywords (comma-separated)")

            keywords = ", ".join(selected_keywords + added_keywords.split(','))

            if content_type == "Blog Post":
                word_count = st.slider("Word Count", min_value=100, max_value=2500, step=100)
            else:
                word_count = st.slider("Word Count", min_value=50, max_value=300, step=10) # Slider for LinkedIn Post

            tone = st.selectbox("Select Tone", ["Casual", "Professional", "Authoritative", "Insightful", "Engaging", "Funny", "Entertaining"])
            reference_urls = st.text_area("Reference URLs (one per line)")
            reference_texts = retrieve_references(reference_urls)

            # Image generation option
            generate_image_option = st.checkbox("Generate Image for the content")
            if generate_image_option:
                image_size = st.selectbox("Select Image Size", ["1024x1024", "1024x1792", "1792x1024"])
                st.session_state.image_size = image_size

            col1, col2 = st.columns(2)

            with col1:
                if st.button("Summon Content"):
                    with st.spinner("Gathering Chakra..."):
                        article = generate_article(client, content_type, title, keywords, tone, word_count, reference_texts, language)
                        st.session_state.article = article
                        if generate_image_option:
                            image_url = generate_image(client, title, st.session_state.image_size)
                            st.session_state.image_url = image_url
                        else:
                            st.session_state.image_url = None

            with col2:
                if st.session_state.image_url and st.button("Re-generate Image"):
                    with st.spinner("Generating new image..."):
                        new_image_url = generate_image(client, title, st.session_state.image_size)
                        st.session_state.image_url = new_image_url
            
            if st.session_state.article:
                st.success("Content summoned successfully!")
                if st.session_state.image_url:
                    st.image(st.session_state.image_url, caption="Generated Image")
                st.markdown(st.session_state.article)
                
                # Add option to download the article as a docx file
                docx_buffer = save_as_docx(title, st.session_state.article, st.session_state.image_url)
                st.download_button(
                    label="Download Article as DOCX",
                    data=docx_buffer,
                    file_name=f"{title}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()